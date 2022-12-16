import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Pt

import SpecData as data
import SysTools
import SysTools as sysTools

"""
spec container: hashmap <String, list>
key: spec number -> String

values:
value 0: div_number -> integer
value 1: div_name -> string
value 2: Name (full name of the spec) -> string
value 3: york_true (if this is an original york region spec) -> boolean
value 4: eto_true (if ETO folder has this spec) -> boolean
value 5: bid_true (should it included in Bid form) -> value
"""
def create_element(name):
    # supporting function adding page
    return OxmlElement(name)


def create_attribute(element, name, value):
    # supporting function adding page
    element.set(ns.qn(name), value)


def add_page_number_odd(paragraph, date):
    # rewrite the entire paragraph that includes page information and date information

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # add second word
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = f'DATE: {date}\t\t'
    of_run._r.append(t2)

    # add first word
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "Page"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)


def add_page_number_even(paragraph, date):
    # rewrite the entire paragraph that includes page information and date information

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # add first word
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "Page"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    # add second word
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = f'\t\t DATE: {date}'
    of_run._r.append(t2)


def update_oddpage_contractno(section, contractNo):
    header = section.header

    # odd page - update contract number
    if len(header.paragraphs) > 0:
        for i in range(0, len(header.paragraphs)):
            if "contr" in header.paragraphs[i].text.lower():
                stringList = header.paragraphs[i].text.split("\t")
                for n in range(0, len(stringList)):
                    if "cont" in stringList[n].lower():
                        header.paragraphs[i].text = header.paragraphs[i].text.replace(stringList[n],
                                                                                      f"CONTRACT NO. {contractNo}")
                        return


def update_evenpage_contractno(section, contractNo):
    header = section.even_page_header

    # even page - update contract number
    if len(header.paragraphs) > 0:
        for i in range(0, len(header.paragraphs)):
            if "contr" in header.paragraphs[i].text.lower():
                stringList = header.paragraphs[i].text.split("\t")
                if len(stringList) > 0:
                    for n in range(0, len(stringList)):
                        if "contr" in stringList[n].lower():
                            header.paragraphs[i].text = header.paragraphs[i].text.replace(stringList[n],
                                                                                          f"CONTRACT NO. {contractNo}")
                            return


def locate_oddpage_date(section):
    paragraphs = section.header.paragraphs

    for i in range(0, len(paragraphs)):
        if "date" in paragraphs[i].text.lower():
            paragraphs[i].text = ""
            return i
    return -1


def locate_evenpage_date(section):
    paragraphs = section.even_page_header.paragraphs

    for i in range(0, len(paragraphs)):
        if "date" in paragraphs[i].text.lower():
            paragraphs[i].text = ""
            return i
    return -1


def update_oddpage_format(section):
    paragraphs = section.header.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri (Body)"
            run.font.size = Pt(11)


def update_evenpage_format(section):
    paragraphs = section.even_page_header.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri (Body)"
            run.font.size = Pt(11)


def update_ETOSpec(path):
    folder_list = sysTools.getFolderNames(path)
    file_dic = sysTools.getFileNames(folder_list)
    sorted_keys = sorted(file_dic.keys())
    file_dic = {key: file_dic[key] for key in sorted_keys}
    result_dic = data.yorkspec_dic

    for division, specs in file_dic.items():
        n = 1
        specs.sort()
        bid_form_code = "A"
        # specs = specs.sort()
        # p
        for spec in specs:
            if len(spec) == 0:
                continue
            if len(spec) > 0 and spec[0].isalpha() == True:
                continue

            word_filepath = division + spec
            key = sysTools.getSpecNumber(spec)

            # initialize bid code number

            if key in result_dic:
                spec_list = result_dic.get(key)
                # value 4
                eto_true = True
                spec_list[4] = eto_true
                # value 5
                bid_true = checkBid(word_filepath)
                if bid_true > 0:
                    bid_form_div_no = spec_list[0]
                    bid_form_code = "A" + bid_form_div_no + '.' + str(n).zfill(2)
                    updateBid(word_filepath, bid_form_code, division, spec)
                    spec_list[5] = bid_form_code
                    n = n + 1
                else:
                    spec_list[5] = "Included but not measured separately"
                result_dic[key] = spec_list
            elif key not in result_dic:
                div_fullname = division.split('/')[-2]
                div_fullname = div_fullname.split(' ')
                # value 0
                div_number = div_fullname[1]
                # value 1
                div_name = data.div_dic.get(div_number)
                # value 2
                spec_name = sysTools.getSpecName(spec, key)
                # value 3
                york_true = False
                # value 4
                eto_true = True
                # value 5
                bid_true = checkBid(word_filepath)
                if bid_true > 0:
                    bid_form_div_no = div_number
                    bid_form_code = "A" + bid_form_div_no + '.' + str(n).zfill(2)
                    updateBid(word_filepath, bid_form_code, division, spec)
                    spec_list[5] = bid_form_code
                    n = n + 1
                else:
                    spec_list[5] = "Included but not measured separately"
                spec_list = [div_number, div_name, spec_name, york_true, eto_true, bid_true]
                result_dic[key] = spec_list
    return result_dic


def checkBid(path):
    file = open(path, 'rb')
    try:
        # avoid file is not word doc.
        document = Document(file)
    except:
        return

    for paragraph in document.paragraphs:
        if "all costs" in paragraph.text.lower() and "bid form" in paragraph.text.lower():
            file.close()
            return 1
    return -1


def updateBid(path, bid_form_code, division, spec):
    bid_form_sentence = "All costs associated with the work of this " \
                        "Section shall be included in the price(s) for " \
                        "Item No. {} in the Bid Form.".format(bid_form_code)

    file = open(path, 'rb')
    try:
        # avoid file is not word doc.
        document = Document(file)
    except:
        return
    target_para = -1
    for i in range(len(document.paragraphs)):
        paragraph = document.paragraphs[i]
        if "all costs" in paragraph.text.lower() and "bid form" in paragraph.text.lower():
            paragraph_list = paragraph.text.split(' ')
            list_len = len(paragraph_list)

            all_loc = -1
            form_loc = -1
            special_case = False

            # locate "all" appear
            for i in range(0, list_len):

                word = paragraph_list[i]
                if 'all' == word.lower():
                    all_loc = i
                if '.1\tall' == word.lower():
                    all_loc = i
                    special_case = True
                if 'form.' == word.lower() or 'form. ' == word.lower() or 'form' == word.lower():
                    form_loc = i
                if all_loc >= 0 and form_loc > 0:
                    for i in range(all_loc, form_loc + 1):
                        paragraph_list[i] = ''
                    if special_case:
                        paragraph_list[all_loc] = '.1\t' + bid_form_sentence
                    else:
                        paragraph_list[all_loc] = bid_form_sentence
                    paragraph.text = ' '.join(paragraph_list)

                    for run in paragraph.runs:
                        run.font.name = "Calibri (Body)"
                        run.font.size = Pt(11)
                    break

    path = division + '/Update/'
    path = SysTools.checkResultPath(path)
    document.save(path + spec)
    file.close()


def getYorkSpec():
    path = "/Users/delinmu/Documents/GitHub/ETO_Specification/YorkOriginal"
    folder_list = sysTools.getFolderNames(path)
    file_dic = sysTools.getFileNames(folder_list)
    sorted_keys = sorted(file_dic.keys())
    file_dic = {key: file_dic[key] for key in sorted_keys}
    spec_dic = {}
    for division, specs in file_dic.items():

        div_fullname = division.split('/')[-2]
        div_fullname = div_fullname.split(' ')
        specs.sort()

        # key
        for spec in specs:
            if len(spec) == 0:
                continue

            if len(spec) > 0 and spec[0].isalpha() == True:
                continue
            key = sysTools.getSpecNumber(spec)
            # value 0
            div_number = div_fullname[1]
            # value 1
            div_name = data.div_dic.get(div_number)
            # value 2
            spec_name = sysTools.getSpecName(spec, key)
            # value 3
            york_true = True
            # value 4
            eto_true = False
            # value 5
            bid_true = False
            spec_list = [div_number, div_name, spec_name, york_true, eto_true, bid_true]
            spec_dic[key] = spec_list

    return spec_dic


def get_ETOSpec_SummarySheets(path, result_path):
    result_dic = update_ETOSpec(path)
    column_title = ['DivisionNumber', 'DivisionName', 'SpecNumber', 'SpecName', 'YorkSpecVersion', 'ETOSpecVersion','BidFormInformation']
    dataframe_list = []
    for key, list in result_dic.items():
        york_true = list[3]
        eto_true = list[4]
        BidFormInformation = list[5]

        if BidFormInformation is False:
            list[5] = "Not Applicable"

        if york_true and eto_true:
            list[3] = "York Region original Spec"
            list[4] = "Included in ETO Spec"

        if york_true is False and eto_true:
            list[3] = "Not Included in York Region original Spec"
            list[4] = "Spec Created by ETO"

        if york_true and eto_true is False:
            list[3] = "York Region original Spec"
            list[4] = "Spec Not Included"

        list.insert(0, key)
        # swap location
        DivisionNumber = list[1]
        DivisionName = list[2]
        SpecNumber = list[0]
        list[0] = DivisionNumber
        list[1] = DivisionName
        list[2] = str(SpecNumber.zfill(6))
        dataframe_list.append(list)
        print(list)

    df = pd.DataFrame(dataframe_list, columns=column_title)

    if (result_path[-1] != '/'):
        result_path = result_path + '/'

    df.to_csv(result_path + "result.csv", index=False)