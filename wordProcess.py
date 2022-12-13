from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns


def create_element(name):
    # supporting function adding page

    return OxmlElement(name)


def create_attribute(element, name, value):
    # supporting function adding page

    element.set(ns.qn(name), value)


def add_page_number_odd(paragraph, date):
    # rewrite the entire paragraph that includes page information and date information

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # add second word
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = f'DATE: {date}\t\t'
    of_run._r.append(t2)

    # add first word
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)


def add_page_number_even(paragraph, date):
    # rewrite the entire paragraph that includes page information and date information

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # add first word
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    # add second word
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = f'\t\t DATE: {date}'
    of_run._r.append(t2)


def update_oddHeader(section, contractNo, dateInfo):
    header = section.header
    # odd page - update contract number
    stringList = header.paragraphs[0].text.split("\t")
    header.paragraphs[0].text = header.paragraphs[0].text.replace(stringList[0], f"CONTRACT NO. {contractNo}")

    # odd page - update date information
    header.paragraphs[2].text = ""


def update_evenHeader(section, contractNo):
    header = section.even_page_header
    # even page - update contract number
    stringList = header.paragraphs[0].text.split("\t")
    header.paragraphs[0].text = header.paragraphs[0].text.replace(stringList[1],
                                                                  f"CONTRACT NO. {contractNo}")
    # even page - remove all content in the page paragraph
    header.paragraphs[2].text = ""


def update_wordInfo(wordPath, contractNo, date):
    file = open(wordPath, 'rb')
    document = Document(file)

    section = document.sections[0]

    # update odd page header
    update_oddHeader(section, contractNo, date)
    add_page_number_odd(section.header.paragraphs[2], date)

    # update even page header
    update_evenHeader(section, contractNo)
    add_page_number_even(section.even_page_header.paragraphs[2], date)

    document.save("your_doc.docx")
    file.close()


wordPath = 'C:/Users/PC/Desktop/Spec/example/02511 Watermains.docx'
contractNo = "SSTR4559"
date = "Oct., 2022"

update_wordInfo(wordPath, contractNo, date)
