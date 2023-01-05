import re

from docx import Document


def get_BidNumber(path):
    file = open(path, 'rb')
    try:
        # avoid file is not word doc.
        document = Document(file)
    except:
        return

    for paragraph in document.paragraphs:
        result = re.search('A[0-1][0-7]\.\d\d', paragraph.text.upper())
        if result:
            return result.group(0)
    return "Missing Bid Number."


def checkBid(path):
    file = open(path, 'rb')
    try:
        # avoid file is not word doc.
        document = Document(file)
    except:
        return

    for paragraph in document.paragraphs:
        if "all costs" in paragraph.text.lower() and "bid form" in paragraph.text.lower():
            file.close()
            return 1
    return -1

def get_ReferenceSpec(path):
    file = open(path, 'rb')
    try:
        # avoid file is not word doc.
        document = Document(file)
    except:
        return
    master = []
    for paragraph in document.paragraphs:
        result = re.findall('Section [0-9]+-[0-9]+|Section [0-9]+|Section [0-9]+[A-Z]', paragraph.text)
        if result:
            for element in result:
                master.append(element)
    return master
