from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Pt


def create_element(name):
    # supporting function adding page
    return OxmlElement(name)


def create_attribute(element, name, value):
    # supporting function adding page
    element.set(ns.qn(name), value)


def add_page_number_odd(paragraph, date):
    # rewrite the entire paragraph that includes page information and date information

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # add second word
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = f'DATE: {date}\t\t'
    of_run._r.append(t2)

    # add first word
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "Page"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)


def add_page_number_even(paragraph, date):
    # rewrite the entire paragraph that includes page information and date information

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # add first word
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "Page"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    # add second word
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = f'\t\t DATE: {date}'
    of_run._r.append(t2)


def update_oddpage_contractno(section, contractNo):
    header = section.header

    # odd page - update contract number
    if len(header.paragraphs) > 0:
        for i in range(0, len(header.paragraphs)):
            if "contr" in header.paragraphs[i].text.lower():
                stringList = header.paragraphs[i].text.split("\t")
                for n in range (0, len(stringList)):
                    if "cont" in stringList[n].lower():
                        header.paragraphs[i].text = header.paragraphs[i].text.replace(stringList[n], f"CONTRACT NO. {contractNo}")
                        return


def update_evenpage_contractno(section, contractNo):
    header = section.even_page_header

    # even page - update contract number
    print(len(header.paragraphs))
    if len(header.paragraphs) > 0:
        for i in range(0, len(header.paragraphs)):
            print(header.paragraphs[i].text)
            if "contr" in header.paragraphs[i].text.lower():
                stringList = header.paragraphs[i].text.split("\t")
                if len(stringList) > 0:
                    for n in range (0, len(stringList)):
                        if "contr" in stringList[n].lower():
                            header.paragraphs[i].text = header.paragraphs[i].text.replace(stringList[n],
                                                                  f"CONTRACT NO. {contractNo}")
                            return


def locate_oddpage_date(section):
    paragraphs = section.header.paragraphs

    for i in range(0, len(paragraphs)):
        if "date" in paragraphs[i].text.lower():
            paragraphs[i].text = ""
            return i
    return -1


def locate_evenpage_date(section):
    paragraphs = section.even_page_header.paragraphs

    for i in range(0, len(paragraphs)):
        if "date" in paragraphs[i].text.lower():
            paragraphs[i].text = ""
            return i
    return -1


def update_oddpage_format(section):
    paragraphs = section.header.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri (Body)"
            run.font.size = Pt(11)

def update_evenpage_format(section):
    paragraphs = section.even_page_header.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri (Body)"
            run.font.size = Pt(11)