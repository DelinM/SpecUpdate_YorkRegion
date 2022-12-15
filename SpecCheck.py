import SysTools as sysTools
import SpecData as data

import docx
from docx import Document
import WordContentUpdate as wordu
import SysTools
from docx.shared import Pt
import pandas as pd

"""
spec container: hashmap <String, list>

key: spec number -> String

values:
value 0: div_number -> integer
value 1: div_name -> string
value 2: Name (full name of the spec) -> string
value 3: york_true (if this is an original york region spec) -> boolean
value 4: eto_true (if ETO folder has this spec) -> boolean
value 5: bid_true (should it included in Bid form) -> value
"""

"""

deleted item
01000
01720
01740
01770
07570
07900
Div 12
Div 14
"""


def getDivName(number):
    return data.div_dic.get(number)


def getYorkSpec():
    path = "/Users/delinmu/Documents/GitHub/ETO_Specification/YorkOriginal"
    folder_list = sysTools.getFolderNames(path)
    file_dic = sysTools.getFileNames(folder_list)
    sorted_keys = sorted(file_dic.keys())
    file_dic = {key: file_dic[key] for key in sorted_keys}
    spec_dic = {}
    for division, specs in file_dic.items():

        div_fullname = division.split('/')[-2]
        div_fullname = div_fullname.split(' ')
        specs.sort()

        # key
        for spec in specs:
            if len(spec) == 0:
                continue

            if len(spec) > 0 and spec[0].isalpha() == True:
                continue
            key = sysTools.getSpecNumber(spec)
            # value 0
            div_number = div_fullname[1]
            # value 1
            div_name = data.div_dic.get(div_number)
            # value 2
            spec_name = sysTools.getSpecName(spec, key)
            # value 3
            york_true = True
            # value 4
            eto_true = False
            # value 5
            bid_true = False
            spec_list = [div_number, div_name, spec_name, york_true, eto_true, bid_true]
            spec_dic[key] = spec_list

    return spec_dic


def update_ETOSpec(path):
    folder_list = sysTools.getFolderNames(path)
    file_dic = sysTools.getFileNames(folder_list)
    sorted_keys = sorted(file_dic.keys())
    file_dic = {key: file_dic[key] for key in sorted_keys}
    result_dic = data.yorkspec_dic

    for division, specs in file_dic.items():
        n = 1
        specs.sort()
        bid_form_code = "A"
        # specs = specs.sort()
        # p
        for spec in specs:
            if len(spec) == 0:
                continue
            if len(spec) > 0 and spec[0].isalpha() == True:
                continue

            word_filepath = division + spec
            key = sysTools.getSpecNumber(spec)

            # initialize bid code number

            if key in result_dic:
                spec_list = result_dic.get(key)
                # value 4
                eto_true = True
                spec_list[4] = eto_true
                # value 5
                bid_true = checkBid(word_filepath)
                if bid_true > 0:
                    bid_form_div_no = spec_list[0]
                    bid_form_code = "A" + bid_form_div_no + '.' + str(n).zfill(2)
                    updateBid(word_filepath, bid_form_code, division, spec)
                    spec_list[5] = bid_form_code
                    n = n + 1
                else:
                    spec_list[5] = "Included but not measured separately"
                result_dic[key] = spec_list
            elif key not in result_dic:
                div_fullname = division.split('/')[-2]
                div_fullname = div_fullname.split(' ')
                # value 0
                div_number = div_fullname[1]
                # value 1
                div_name = data.div_dic.get(div_number)
                # value 2
                spec_name = sysTools.getSpecName(spec, key)
                # value 3
                york_true = False
                # value 4
                eto_true = True
                # value 5
                bid_true = checkBid(word_filepath)
                if bid_true > 0:
                    bid_form_div_no = div_number
                    bid_form_code = "A" + bid_form_div_no + '.' + str(n).zfill(2)
                    updateBid(word_filepath, bid_form_code, division, spec)
                    spec_list[5] = bid_form_code
                    n = n + 1
                else:
                    spec_list[5] = "Included but not measured separately"
                spec_list = [div_number, div_name, spec_name, york_true, eto_true, bid_true]
                result_dic[key] = spec_list
    print("run")
    return result_dic


def checkBid(path):
    file = open(path, 'rb')
    try:
        # avoid file is not word doc.
        document = Document(file)
    except:
        return

    for paragraph in document.paragraphs:
        if "all costs" in paragraph.text.lower() and "bid form" in paragraph.text.lower():
            file.close()
            return 1
    return -1


def updateBid(path, bid_form_code, division, spec):
    bid_form_sentence = "All costs associated with the work of this " \
                        "Section shall be included in the price(s) for " \
                        "Item No. {} in the Bid Form.".format(bid_form_code)

    file = open(path, 'rb')
    try:
        # avoid file is not word doc.
        document = Document(file)
    except:
        return
    target_para = -1
    for i in range(len(document.paragraphs)):
        paragraph = document.paragraphs[i]
        if "all costs" in paragraph.text.lower() and "bid form" in paragraph.text.lower():
            paragraph_list = paragraph.text.split(' ')
            list_len = len(paragraph_list)

            all_loc = -1
            form_loc = -1
            special_case = False

            # locate "all" appear
            for i in range(0, list_len):

                word = paragraph_list[i]
                if 'all' == word.lower():
                    all_loc = i
                if '.1\tall' == word.lower():
                    all_loc = i
                    special_case = True
                if 'form.' == word.lower() or 'form. ' == word.lower() or 'form' == word.lower():
                    form_loc = i
                if all_loc >= 0 and form_loc > 0:
                    for i in range(all_loc, form_loc + 1):
                        paragraph_list[i] = ''
                    if special_case:
                        paragraph_list[all_loc] = '.1\t' + bid_form_sentence
                    else:
                        paragraph_list[all_loc] = bid_form_sentence
                    paragraph.text = ' '.join(paragraph_list)

                    for run in paragraph.runs:
                        run.font.name = "Calibri (Body)"
                        run.font.size = Pt(11)
                    break

    path = division + '/Update/'
    path = SysTools.checkResultPath(path)
    document.save(path + spec)
    file.close()


"""
spec container: hashmap <String, list>

key: spec number -> String

values:
value 0: div_number -> integer
value 1: div_name -> string
value 2: Name (full name of the spec) -> string
value 3: york_true (if this is an original york region spec) -> boolean
value 4: eto_true (if ETO folder has this spec) -> boolean
value 5: bid_true (should it included in Bid form) -> value
"""


def get_ETOSpec_SummarySheets(path):
    result_dic = update_ETOSpec(path)
    #                       0               1               2           3               4               5               6
    column_title = ['DivisionNumber', 'DivisionName', 'SpecNumber', 'SpecName', 'YorkSpecVersion', 'ETOSpecVersion','BidFormInformation']
    dataframe_list = []
    for key, list in result_dic.items():
        york_true = list[3]
        eto_true = list[4]
        BidFormInformation = list[5]

        if BidFormInformation is False:
            list[5] = "Not Applicable"

        if york_true and eto_true:
            list[3] = "York Region original Spec"
            list[4] = "Included in ETO Spec"

        if york_true is False and eto_true:
            list[3] = "Not Included in York Region original Spec"
            list[4] = "Spec Created by ETO"

        if york_true and eto_true is False:
            list[3] = "York Region original Spec"
            list[4] = "Spec Not Included"

        list.insert(0, key)
        # swap location
        DivisionNumber = list[1]
        DivisionName = list[2]
        SpecNumber = list[0]
        list[0] = DivisionNumber
        list[1] = DivisionName
        list[2] = str(SpecNumber.zfill(6))
        dataframe_list.append(list)
        print(list[2])

    df = pd.DataFrame(dataframe_list, columns=column_title)
    df = df.astype(str)

    df.to_csv("result.csv")
    print(df)


get_ETOSpec_SummarySheets('/Users/delinmu/Documents/GitHub/ETO_Specification/Spec')
