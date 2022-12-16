import docx
from docx import Document

import SpecData as data
import WordContent as wordu


def getWordPaths(dic):
    word_path = []

    for key, values in dic.items():
        for value in values:
            word_path.append(key + value)
    return word_path


def update_wordInfo(wordPath, contractNo, date, name):
    file = open(wordPath, 'rb')
    document = Document(file)

    section = document.sections[0]

    # update odd page header
    wordu.update_oddpage_contractno(section, contractNo)
    date_odd_location = wordu.locate_oddpage_date(section)
    if date_odd_location > 0:
        wordu.add_page_number_odd(section.header.paragraphs[date_odd_location], date)
        wordu.update_oddpage_format(section)
    else:
        document.save(name)
        return


    # update even page header
    wordu.update_evenpage_contractno(section, contractNo)
    date_even_location = wordu.locate_evenpage_date(section)
    if date_even_location > 0:
        wordu.add_page_number_even(section.even_page_header.paragraphs[date_even_location], date)
        wordu.update_evenpage_format(section)


    document.save(name)
    file.close()


def getDivName(number):
    return data.div_dic.get(number)
